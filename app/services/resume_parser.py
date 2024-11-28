import os
import logging
from typing import Dict, Optional, List
import PyPDF2
from docx import Document
import re
import time
import pdfplumber
import docx2txt
from PIL import Image
import pytesseract
import io
import numpy as np
from pdf2image import convert_from_path

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx'}
        self.timeout = 60  # 60 seconds timeout
        self.max_file_size = 16 * 1024 * 1024  # 16MB
        self.education_keywords = [
            'education', 'academic', 'qualification', 'degree', 'university', 'college',
            'học vấn', 'trường', 'đại học', 'cao đẳng', 'bằng cấp', 'chứng chỉ'
        ]
        self.experience_keywords = [
            'experience', 'employment', 'work history', 'professional background',
            'kinh nghiệm', 'công việc', 'dự án', 'làm việc', 'chức vụ'
        ]
        self.skills_keywords = [
            'skills', 'technical skills', 'competencies', 'expertise',
            'kỹ năng', 'chuyên môn', 'công nghệ', 'ngôn ngữ'
        ]

    def parse(self, file_path: str) -> Dict:
        """Parse resume file and extract information"""
        start_time = time.time()
        logger.info(f"Starting to parse resume: {file_path}")
        
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
                
            _, ext = os.path.splitext(file_path)
            if ext.lower() not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {ext}")

            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                raise ValueError(f"File too large (max {self.max_file_size/(1024*1024)}MB)")

            # Try multiple methods to extract text
            text = ""
            if ext.lower() == '.pdf':
                text = self._parse_pdf_advanced(file_path)
            else:
                text = self._parse_docx_advanced(file_path)
                
            # Check timeout
            if time.time() - start_time > self.timeout:
                raise TimeoutError(f"Parsing timeout after {self.timeout} seconds")
                
            # Extract information
            result = self._extract_information(text)
            logger.info(f"Resume parsing completed in {time.time() - start_time:.2f} seconds")
            return result

        except Exception as e:
            logger.exception(f"Error parsing resume: {str(e)}")
            return {
                'error': str(e),
                'name': 'Unknown',
                'email': '',
                'phone': '',
                'education': '',
                'experience': '',
                'skills': [],
                'raw_text': ''
            }

    def _parse_pdf_advanced(self, file_path: str) -> str:
        """Parse PDF using multiple methods for better text extraction"""
        logger.info(f"Starting advanced PDF parsing: {file_path}")
        text = ""
        
        # Method 1: Try PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")

        # If PyPDF2 failed or extracted no text, try pdfplumber
        if not text.strip():
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {str(e)}")

        # If still no text, try OCR
        if not text.strip():
            try:
                images = convert_from_path(file_path)
                for image in images:
                    text += pytesseract.image_to_string(image, lang='vie+eng') + "\n"
            except Exception as e:
                logger.warning(f"OCR extraction failed: {str(e)}")

        if not text.strip():
            raise ValueError("Could not extract text from PDF using any method")

        return text

    def _parse_docx_advanced(self, file_path: str) -> str:
        """Parse DOCX using multiple methods for better text extraction"""
        logger.info(f"Starting advanced DOCX parsing: {file_path}")
        text = ""
        
        # Method 1: Try python-docx
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {str(e)}")

        # Method 2: Try docx2txt if python-docx failed
        if not text.strip():
            try:
                text = docx2txt.process(file_path)
            except Exception as e:
                logger.warning(f"docx2txt extraction failed: {str(e)}")

        if not text.strip():
            raise ValueError("Could not extract text from DOCX using any method")

        return text

    def _extract_information(self, text: str) -> Dict:
        """Extract structured information from text"""
        logger.info("Starting information extraction")
        
        # Normalize text
        text = self._normalize_text(text)
        
        # Extract basic information
        email = self._extract_email(text)
        phone = self._extract_phone(text)
        name = self._extract_name(text)
        
        # Extract sections
        education = self._extract_section(text, self.education_keywords)
        experience = self._extract_section(text, self.experience_keywords)
        skills = self._extract_skills(text)

        return {
            'name': name,
            'email': email,
            'phone': phone,
            'education': education,
            'experience': experience,
            'skills': skills,
            'raw_text': text
        }

    def _normalize_text(self, text: str) -> str:
        """Normalize text for better extraction"""
        # Replace multiple newlines with single newline
        text = re.sub(r'\n+', '\n', text)
        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters
        text = re.sub(r'[^\w\s@.+-]', ' ', text)
        return text.strip()

    def _extract_email(self, text: str) -> str:
        """Extract email from text"""
        patterns = [
            r'[\w\.-]+@[\w\.-]+\.\w+',  # Basic email pattern
            r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+' # More specific pattern
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return ''

    def _extract_phone(self, text: str) -> str:
        """Extract phone number from text"""
        patterns = [
            r'(?:\+84|0)\d{9,10}',  # Vietnamese numbers
            r'\d{3}[-.]?\d{3}[-.]?\d{4}',  # General format
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}'  # (123) 456-7890
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return ''

    def _extract_name(self, text: str) -> str:
        """Extract name from text using improved patterns"""
        # Common Vietnamese family names
        family_names = ['Nguyễn', 'Trần', 'Lê', 'Phạm', 'Hoàng', 'Huỳnh', 'Phan', 'Vũ', 'Võ', 'Đặng', 'Bùi', 'Đỗ']
        
        # Try to find name at the beginning of the document
        lines = text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            
            # Skip lines that are clearly not names
            if not line or any(w in line.lower() for w in ['resume', 'cv', 'curriculum']):
                continue
                
            # Check for Vietnamese names
            if any(name in line for name in family_names):
                return line
                
            # Check for Vietnamese diacritical marks
            if any(c in line for c in 'áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ'):
                words = line.split()
                if 2 <= len(words) <= 6:  # Most Vietnamese names are 2-6 words
                    return line
        
        return 'Unknown'

    def _extract_section(self, text: str, keywords: List[str]) -> str:
        """Extract a section from text based on keywords"""
        text_lower = text.lower()
        
        # Find the start of the section
        start_pos = -1
        start_keyword = ''
        for keyword in keywords:
            pos = text_lower.find(keyword.lower())
            if pos != -1 and (start_pos == -1 or pos < start_pos):
                start_pos = pos
                start_keyword = keyword

        if start_pos == -1:
            return ''

        # Find the end of the section (next section or end of text)
        end_pos = len(text)
        for keyword in self.education_keywords + self.experience_keywords + self.skills_keywords:
            if keyword.lower() != start_keyword.lower():
                pos = text_lower.find(keyword.lower(), start_pos + len(start_keyword))
                if pos != -1 and pos < end_pos:
                    end_pos = pos

        section_text = text[start_pos:end_pos].strip()
        # Remove the section header
        section_text = re.sub(f'^{start_keyword}[:\s]*', '', section_text, flags=re.IGNORECASE)
        return section_text.strip()

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        # Get skills section
        skills_text = self._extract_section(text, self.skills_keywords)
        if not skills_text:
            return []

        # Split into individual skills
        skills = []
        # Split by common separators
        for skill in re.split(r'[,•\n]', skills_text):
            skill = skill.strip()
            if skill and len(skill) > 1:  # Ignore single characters
                skills.append(skill)

        return list(set(skills))  # Remove duplicates
