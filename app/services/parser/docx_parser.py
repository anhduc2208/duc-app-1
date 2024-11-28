from docx import Document
import re

class DocxParser:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)
    
    def extract_text(self):
        """Extract all text from DOCX file"""
        return '\n'.join([paragraph.text for paragraph in self.document.paragraphs])
    
    def extract_info(self):
        """Extract basic information from resume"""
        text = self.extract_text()
        
        # Find email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email = re.findall(email_pattern, text)
        
        # Find phone number
        phone_pattern = r'(\+84|0)[0-9]{9,10}'
        phone = re.findall(phone_pattern, text)
        
        # Find name (assume it's in the first line)
        name = text.split('\n')[0].strip()
        
        return {
            'name': name,
            'email': email[0] if email else None,
            'phone': phone[0] if phone else None
        }
    
    def extract_sections(self):
        """Extract different sections from resume"""
        text = self.extract_text()
        sections = {
            'education': '',
            'experience': '',
            'skills': []
        }
        
        paragraphs = [p.text.strip() for p in self.document.paragraphs if p.text.strip()]
        current_section = None
        section_text = []
        
        for paragraph in paragraphs:
            # Detect section headers
            paragraph_lower = paragraph.lower()
            if 'education' in paragraph_lower or 'học vấn' in paragraph_lower:
                if current_section:
                    sections[current_section] = '\n'.join(section_text)
                current_section = 'education'
                section_text = []
            elif 'experience' in paragraph_lower or 'kinh nghiệm' in paragraph_lower:
                if current_section:
                    sections[current_section] = '\n'.join(section_text)
                current_section = 'experience'
                section_text = []
            elif 'skills' in paragraph_lower or 'kỹ năng' in paragraph_lower:
                if current_section:
                    sections[current_section] = '\n'.join(section_text)
                current_section = 'skills'
                section_text = []
            elif current_section:
                section_text.append(paragraph)
        
        # Add the last section
        if current_section and section_text:
            if current_section == 'skills':
                sections[current_section] = [skill.strip() for skill in section_text]
            else:
                sections[current_section] = '\n'.join(section_text)
        
        # Extract skills from skills section if it's a string
        if isinstance(sections['skills'], str):
            sections['skills'] = [skill.strip() for skill in sections['skills'].split(',')]
        
        return sections
